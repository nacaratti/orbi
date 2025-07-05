import discord
from discord.ext import commands
import asyncio
import os
import datetime
import tempfile
from openai import OpenAI
import docx
from pydub import AudioSegment
from dotenv import load_dotenv
import nacl

load_dotenv()

OPENAI_API_KEY= os.getenv("OPENAI_API_KEY")
DISCORD_BOT_TOKEN= os.getenv("DISCORD_BOT_TOKEN")

intents = discord.Intents.default()
intents.message_content = True
intents.voice_states = True
bot = commands.Bot(command_prefix="!", intents=discord.Intents.all())

client = OpenAI(api_key = OPENAI_API_KEY)

active_recordings = {}

class VoiceRecorder:
    def __init__(self, voice_channel, text_channel):
        self.voice_channel = voice_channel
        self.text_channel = text_channel
        self.voice_client = None
        self.recording = False
        self.combined_audio = None
        self.user_audio = {}
        self.start_time = None
        self.temp_dir = tempfile.mkdtemp()

    async def start_recording(self):
        self.voice_client = await self.voice_channel.connect()
        self.recording = True
        self.start_time = datetime.datetime.now()
        self.combined_audio = AudioSegment.silent(duration=0)

        await self.text_channel.send(f"📝 Iniciando gravação na sala {self.voice_channel.name}")

        self.sink = discord.sinks.MP3Sink()
        self.voice_client.start_recording(
            self.sink,
            self.on_recording_finished,
            self.voice_channel
        )

    async def stop_recording(self):
        if self.voice_client and self.voice_client.recording:
            self.voice_client.stop_recording()
            self.recording = False
            await self.voice_client.disconnect()
            await self.text_channel.send("✅ Gravação finalizada. Processando áudio...")
        else:
            await self.text_channel.send("❌ Não há gravação ativa para interromper.")

    async def on_recording_finished(self, sink, channel, *args):
        combined_file = os.path.join(self.temp_dir, f"combined_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3")

        all_audio = AudioSegment.silent(duration=0)
        for user_id, audio in sink.audio_data.items():
            user_file = os.path.join(self.temp_dir, f"user_{user_id}.mp3")
            with open(user_file, "wb") as f:
                f.write(audio.file.read())
            user_audio = AudioSegment.from_mp3(user_file)
            if len(all_audio) < len(user_audio):
                all_audio = all_audio + AudioSegment.silent(duration=len(user_audio) - len(all_audio))
            all_audio = all_audio.overlay(user_audio)

        all_audio.export(combined_file, format="mp3")
        await self.text_channel.send("🔄 Áudio combinado gerado. Transcrevendo...")

        await self.transcribe_and_generate_summary(combined_file)

    async def transcribe_and_generate_summary(self, audio_file):
        try:
            await self.text_channel.send("🔄 Transcrevendo áudio...")
            with open(audio_file, "rb") as file:
                transcription = await asyncio.to_thread(
                    client.audio.transcriptions.create,
                    file=file,
                    model="whisper-1"
                )
            transcript_text = transcription.text

            await self.text_channel.send("🔄 Gerando ata da reunião...")

            response = await asyncio.to_thread(
                client.chat.completions.create,
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Você é um assistente especializado em gerar atas de reunião a partir de transcrições. Crie uma ata formal com: 1) Data e horário, 2) Participantes (usuários do Discord), 3) Pontos discutidos, 4) Decisões, 5) Próximos passos."},
                    {"role": "user", "content": f"Transcrição: {transcript_text}"}
                ]
            )
            summary = response.choices[0].message.content

            doc = docx.Document()
            doc_title = f"Ata de Reunião - {self.start_time.strftime('%d/%m/%Y %H:%M')}"
            doc.add_heading(doc_title, 0)
            doc.add_paragraph(summary)

            doc_file = os.path.join(self.temp_dir, f"ata_reuniao_{self.start_time.strftime('%Y%m%d_%H%M')}.docx")
            doc.save(doc_file)

            await self.text_channel.send("🎧 Áudio gravado:", file=discord.File(audio_file))
            await self.text_channel.send("📄 Ata da reunião gerada:", file=discord.File(doc_file))

            os.remove(audio_file)
            os.remove(doc_file)

        except Exception as e:
            await self.text_channel.send(f"❌ Erro ao processar o áudio: {str(e)}")

@bot.event
async def on_ready():
    print(f"✅ Bot online como {bot.user}!")

@bot.slash_command(name="gravar", description="Inicia a gravação na sala de voz")
async def gravar(ctx: discord.ApplicationContext):
    if ctx.author.voice is None:
        await ctx.respond("❌ Você precisa estar em um canal de voz.", ephemeral=True)
        return

    await ctx.defer()  # Defer para ganhar mais tempo
    
    voice_channel = ctx.author.voice.channel
    if voice_channel.id in active_recordings:
        await ctx.respond("❌ Já existe uma gravação ativa neste canal.", ephemeral=True)
        return

    recorder = VoiceRecorder(voice_channel, ctx.channel)
    active_recordings[voice_channel.id] = recorder
    await recorder.start_recording()

    # Aqui, como já deferiu, use followup
    await ctx.followup.send(f"🔴 Gravação iniciada na sala {voice_channel.name}.")

@bot.slash_command(name="parar", description="Para a gravação na sala de voz")
async def parar(ctx: discord.ApplicationContext):
    if ctx.author.voice is None:
        await ctx.respond("❌ Você precisa estar em um canal de voz.", ephemeral=True)
        return

    voice_channel = ctx.author.voice.channel
    if voice_channel.id not in active_recordings:
        await ctx.respond("❌ Não há gravação ativa neste canal.", ephemeral=True)
        return

    await ctx.defer()  # sinaliza que vai demorar um pouco para responder

    recorder = active_recordings[voice_channel.id]
    await recorder.stop_recording()
    del active_recordings[voice_channel.id]

    await ctx.followup.send("🛑 Gravação encerrada e sendo processada.")


@bot.slash_command(name="ajuda", description="Mostra os comandos do bot")
async def ajuda(ctx: discord.ApplicationContext):
    embed = discord.Embed(title="Comandos disponíveis", color=discord.Color.blurple())
    embed.add_field(name="/gravar", value="Inicia a gravação de áudio", inline=False)
    embed.add_field(name="/parar", value="Finaliza a gravação de áudio", inline=False)
    embed.add_field(name="/ajuda", value="Exibe esta mensagem", inline=False)
    await ctx.respond(embed=embed)

# Rodar o bot
bot.run(DISCORD_BOT_TOKEN)